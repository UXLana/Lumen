#!/usr/bin/env python3
"""Generate VPAT Accessibility Remediation Report as .docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

# Colors
DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
MEDIUM_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
CRITICAL_RED = RGBColor(0xC0, 0x39, 0x2B)
SERIOUS_ORANGE = RGBColor(0xD4, 0x6B, 0x08)
PASS_GREEN = RGBColor(0x27, 0xAE, 0x60)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
HEADER_BG = RGBColor(0x1B, 0x3A, 0x5C)

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        set_cell_shading(cell, '1B3A5C')

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F5F5F5')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table

def add_heading_styled(doc, text, level=1):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE
    return h

def add_ticket_header(doc, ticket_id, title, priority, issues_fixed, wcag, component, est_time):
    """Add a formatted ticket header block."""
    doc.add_paragraph()  # spacer

    # Ticket title bar
    p = doc.add_paragraph()
    run = p.add_run(f'{ticket_id}: {title}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE

    # Metadata table
    table = doc.add_table(rows=2, cols=4)
    table.autofit = True

    labels = ['Priority', 'Issues Fixed', 'WCAG', 'Est. Time']
    values = [priority, str(issues_fixed), wcag, est_time]

    for i, label in enumerate(labels):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, '2C5F8A')

    for i, value in enumerate(values):
        cell = table.rows[1].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(9)
        if 'CRITICAL' in priority:
            if i == 0:
                run.font.color.rgb = CRITICAL_RED
                run.bold = True
        elif 'SERIOUS' in priority:
            if i == 0:
                run.font.color.rgb = SERIOUS_ORANGE
                run.bold = True

    # Component line
    p = doc.add_paragraph()
    run = p.add_run('Component: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(component)
    run.font.size = Pt(9)

def add_code_block(doc, code_text):
    """Add a monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

def add_checklist(doc, items):
    """Add verification checklist items."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'\u2610  {item}')
        run.font.size = Pt(9)

def add_section_label(doc, text):
    """Add a bold section label within a ticket."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = MEDIUM_BLUE


# ============================================================
# DOCUMENT CONTENT
# ============================================================

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('VPAT Accessibility Remediation')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Same-Day Fix Plan')
run.font.size = Pt(16)
run.font.color.rgb = MEDIUM_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Date: February 12, 2026  |  Product: Lumen v1.17.19225  |  VPAT Date: 1/13/2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # spacer

# ============================================================
# VPAT CONFORMANCE SUMMARY
# ============================================================
add_heading_styled(doc, 'VPAT Conformance Summary', level=1)

p = doc.add_paragraph()
run = p.add_run('The VPAT report (dated 1/13/2026) identified ')
run.font.size = Pt(10)
run = p.add_run('5 criteria')
run.bold = True
run.font.size = Pt(10)
run = p.add_run(' with "Partially Supports" status:')
run.font.size = Pt(10)

add_styled_table(doc,
    ['WCAG Criterion', 'VPAT Status', 'Remediation'],
    [
        ['1.3.1 Info and Relationships (A)', 'Partially Supports', 'A11Y-001, A11Y-002, A11Y-003'],
        ['2.5.2 Pointer Cancellation (A)', 'Partially Supports', 'A11Y-006'],
        ['4.1.2 Name, Role, Value (A)', 'Partially Supports', 'A11Y-004'],
        ['1.4.3 Contrast Minimum (AA)', 'Partially Supports', 'A11Y-005'],
        ['2.4.5 Multiple Ways (AA)', 'Partially Supports', 'Out of scope*'],
    ]
)

p = doc.add_paragraph()
run = p.add_run('*2.4.5 relates to training videos hosted externally \u2014 content/UX issue, not a code fix.')
run.font.size = Pt(8)
run.font.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
add_heading_styled(doc, 'Executive Summary', level=1)

add_styled_table(doc,
    ['Ticket', 'Component', 'Issues', 'Severity', 'Time', 'WCAG'],
    [
        ['A11Y-001', 'Main Navigation Bar', '157', 'Critical', '1-2 hrs', '1.3.1'],
        ['A11Y-002', 'Tabstrip Components', '11', 'Critical', '30 min', '1.3.1'],
        ['A11Y-003', 'Nav + Tab List Structure', '102', 'Serious', '\u2014', '(auto)'],
        ['A11Y-004', 'Grid ARIA Attributes', '6', 'Critical', '30 min', '4.1.2'],
        ['A11Y-005', 'Color Contrast', '12+', 'Serious', '30 min', '1.4.3'],
        ['A11Y-006', 'Pointer Cancellation', 'TBD', 'Serious', '1 hr', '2.5.2'],
        ['A11Y-007', 'External Training Content', 'N/A', 'Low', 'N/A', '2.4.5'],
    ]
)

p = doc.add_paragraph()
run = p.add_run('Total Estimated Time: 4-5 hours of dev work')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = DARK_BLUE

# Issue distribution
add_heading_styled(doc, 'Issue Distribution by Page (axe scan)', level=2)

add_styled_table(doc,
    ['Page', 'Critical', 'Serious', 'Total'],
    [
        ['Plants', '39', '15', '54'],
        ['Packages', '27', '27', '54'],
        ['Transfers', '26', '15', '41'],
        ['Financials', '20', '15', '35'],
        ['Reports', '19', '14', '33'],
        ['Admin', '23', '15', '38'],
        ['User Profile', '19', '14', '33'],
    ]
)

# Root cause
add_heading_styled(doc, 'Root Cause Analysis', level=2)
p = doc.add_paragraph()
run = p.add_run('87% of all issues (252 of 288)')
run.bold = True
run.font.size = Pt(10)
run = p.add_run(' stem from just 2 root causes in shared UI components:')
run.font.size = Pt(10)

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('Navigation bar ARIA hierarchy')
run.bold = True
run.font.size = Pt(10)
run = p.add_run(' \u2014 <li> elements break the menubar > menuitem parent-child relationship (shared across all 7 pages)')
run.font.size = Pt(10)

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('Kendo tabstrip ARIA structure')
run.bold = True
run.font.size = Pt(10)
run = p.add_run(' \u2014 role="tablist" is on a <div> wrapper instead of the <ul>, so tab <li> children lose their parent context')
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Fixing these 2 root components resolves ~252 of 288 issues automatically.')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = PASS_GREEN

# ============================================================
# PAGE BREAK BEFORE TICKETS
# ============================================================
doc.add_page_break()

add_heading_styled(doc, 'Remediation Tickets', level=1)

# ============================================================
# A11Y-001
# ============================================================
add_ticket_header(doc,
    'A11Y-001', 'Fix Main Navigation Bar ARIA Parent-Child Hierarchy',
    'P0 - CRITICAL', 157, '1.3.1 Info and Relationships (Level A)',
    'Main Navigation Bar (ul.nav.d-flex[role="menubar"])', '1-2 hours'
)

add_section_label(doc, 'VPAT Finding')
p = doc.add_paragraph()
run = p.add_run('"The main navigation bar and grids that have tabs include ARIA roles that do not have the appropriate required parent/children relationships"')
run.font.size = Pt(9)
run.italic = True

add_section_label(doc, 'Problem')
p = doc.add_paragraph()
run = p.add_run('All <a> elements with role="menuitem" are not contained within a proper role="menu" or role="menubar" parent. The <li> elements break the expected ARIA parent-child relationship. menubar can only contain menuitem, menuitemcheckbox, menuitemradio, or elements with role="none"/"presentation".')
run.font.size = Pt(9)

p = doc.add_paragraph()
run = p.add_run('Axe reports 2 distinct rule violations from this single root cause:')
run.font.size = Pt(9)

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('aria-required-parent: ')
run.bold = True
run.font.size = Pt(9)
run = p.add_run('18 menuitem links per page lack a valid menubar/menu parent (126 total across 7 pages)')
run.font.size = Pt(9)

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('aria-required-children: ')
run.bold = True
run.font.size = Pt(9)
run = p.add_run('1 instance per page where the menubar has invalid <li> children (7 total)')
run.font.size = Pt(9)

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('listitem: ')
run.bold = True
run.font.size = Pt(9)
run = p.add_run('14 orphaned <li> elements per page (98 total, auto-resolved \u2014 see A11Y-003)')
run.font.size = Pt(9)

add_section_label(doc, 'Affected Elements (per page)')
elements = [
    '.dropdown:nth-child(1) \u2014 Notifications dropdown',
    '.facilities-dropdown \u2014 Facilities selector',
    '.navigation-first-link \u2014 First nav item (contextual)',
    '.split-dropdown.dropdown (x5) \u2014 Plants, Packages, Transfers, Financials, Reports, Admin',
    '.separator \u2014 Flex spacer',
    '.search-menu \u2014 Search',
    '.import-menu \u2014 Data Import',
    '.support-menu \u2014 Support',
    '.sessionstate-menu \u2014 Session/Saved',
    '.user-menu \u2014 User profile',
]
for el in elements:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(el)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'

add_section_label(doc, 'Fix (Recommended)')
p = doc.add_paragraph()
run = p.add_run('Add role="none" to all <li> elements that are direct children of ul[role="menubar"]. This makes them presentational, restoring the valid menubar > menuitem hierarchy.')
run.font.size = Pt(9)

add_code_block(doc, '<!-- Before -->\n<li class="dropdown split-dropdown">\n  <a role="menuitem" ...>Plants</a>\n</li>')
add_code_block(doc, '<!-- After -->\n<li class="dropdown split-dropdown" role="none">\n  <a role="menuitem" ...>Plants</a>\n</li>')

add_section_label(doc, 'Additional Fix: Data Import Link')
p = doc.add_paragraph()
run = p.add_run('Change aria-label="null" to aria-label="Data Import" on the Data Import menu item.')
run.font.size = Pt(9)

add_code_block(doc, '<!-- Before -->\n<a href="...dataimport/queue" role="menuitem" aria-label="null" alt="null">')
add_code_block(doc, '<!-- After -->\n<a href="...dataimport/queue" role="menuitem" aria-label="Data Import" alt="Data Import">')

add_section_label(doc, 'Verification')
add_checklist(doc, [
    'axe scan shows 0 aria-required-parent errors on menuitem elements',
    'axe scan shows 0 aria-required-children errors on menubar',
    'Keyboard: Tab navigates through all menu items in order',
    'Screen reader: Items announced as "menuitem" within "menubar"',
    'Run axe on all 7 pages: Plants, Packages, Transfers, Financials, Reports, Admin, User Profile',
    'VPAT 1.3.1 can be updated to "Supports"',
])

# ============================================================
# A11Y-002
# ============================================================
doc.add_page_break()

add_ticket_header(doc,
    'A11Y-002', 'Fix Kendo Tabstrip ARIA Role Placement',
    'P0 - CRITICAL', 11, '1.3.1 Info and Relationships (Level A)',
    'Kendo UI Tabstrip (k-tabstrip widgets)', '30 minutes'
)

add_section_label(doc, 'VPAT Finding')
p = doc.add_paragraph()
run = p.add_run('"The main navigation bar and grids that have tabs include ARIA roles that do not have the appropriate required parent/children relationships"')
run.font.size = Pt(9)
run.italic = True

add_section_label(doc, 'Problem')
p = doc.add_paragraph()
run = p.add_run('Kendo tabstrip components render with role="tablist" on the outer <div>, but the actual tab <li role="tab"> elements are inside a <ul> child. This creates two problems:')
run.font.size = Pt(9)

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('The <div role="tablist"> has an invalid child <ul> (should contain tab elements directly)')
run.font.size = Pt(9)

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('The <li role="tab"> elements don\'t have a direct tablist parent')
run.font.size = Pt(9)

add_section_label(doc, 'Affected Components')
components = [
    '#plants_tabstrip \u2014 Plants page (19 tabs)',
    '#packages_tabstrip \u2014 Packages page',
    '#transfers_tabstrip \u2014 Transfers page',
    '#tagorders_tabstrip \u2014 Admin page',
]
for c in components:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(c)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'

add_section_label(doc, 'Fix')
p = doc.add_paragraph()
run = p.add_run('Move role="tablist" from the outer <div> to the <ul.k-tabstrip-items> element:')
run.font.size = Pt(9)

add_code_block(doc, '// Apply after Kendo renders tabstrips\n$(\'.k-tabstrip\').removeAttr(\'role\');\n$(\'.k-tabstrip-items\').attr(\'role\', \'tablist\');')

add_code_block(doc, '<!-- Before -->\n<div id="plants_tabstrip" role="tablist">\n  <ul class="k-tabstrip-items">\n    <li role="tab">...</li>\n  </ul>\n</div>')

add_code_block(doc, '<!-- After -->\n<div id="plants_tabstrip">\n  <ul class="k-tabstrip-items" role="tablist">\n    <li role="tab">...</li>\n  </ul>\n</div>')

add_section_label(doc, 'Verification')
add_checklist(doc, [
    'axe scan shows 0 aria-required-children errors on tabstrip containers',
    'axe scan shows 0 aria-required-parent errors on tab elements',
    'Keyboard: Left/Right arrow keys navigate between tabs',
    'Screen reader: Tabs announced as "tab 1 of N" with correct labeling',
    'VPAT 1.3.1 can be updated to "Supports"',
])

# ============================================================
# A11Y-003
# ============================================================
doc.add_page_break()

add_ticket_header(doc,
    'A11Y-003', 'Resolve Orphaned List Item Structure (Auto-Fix)',
    'P1 - SERIOUS', 102, '1.3.1 Info and Relationships (Level A)',
    'Navigation <li> elements and tabstrip <li> elements', '0 (resolved by A11Y-001 and A11Y-002)'
)

add_section_label(doc, 'Description')
p = doc.add_paragraph()
run = p.add_run('These issues are ')
run.font.size = Pt(9)
run = p.add_run('automatically resolved')
run.bold = True
run.font.size = Pt(9)
run = p.add_run(' by A11Y-001 and A11Y-002. No additional code changes required.')
run.font.size = Pt(9)

p = doc.add_paragraph()
run = p.add_run('Axe flags listitem and list violations because:')
run.font.size = Pt(9)

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('Navigation <li> elements are inside a <ul role="menubar">, which overrides list semantics. When role="none" is applied to <li> elements (A11Y-001), they become presentational and are no longer flagged. (98 issues)')
run.font.size = Pt(9)

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('Tabstrip <ul> contains <li role="tab"> elements, making them invalid list children. When role="tablist" moves to <ul> (A11Y-002), the tab structure becomes valid. (4 issues)')
run.font.size = Pt(9)

p = doc.add_paragraph()
run = p.add_run('Blocked by: A11Y-001, A11Y-002')
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = SERIOUS_ORANGE

add_section_label(doc, 'Verification')
add_checklist(doc, [
    'axe scan shows 0 listitem errors across all pages',
    'axe scan shows 0 list errors across all pages',
])

# ============================================================
# A11Y-004
# ============================================================
add_ticket_header(doc,
    'A11Y-004', 'Fix Invalid ARIA Attribute Values on Grid Rows',
    'P0/P1 - CRITICAL', 6, '4.1.2 Name, Role, Value (Level A)',
    'Kendo Grid <tr> elements', '30 minutes'
)

add_section_label(doc, 'VPAT Finding')
p = doc.add_paragraph()
run = p.add_run('"Some ARIA attributes do not conform to valid values"')
run.font.size = Pt(9)
run.italic = True

add_section_label(doc, 'Problem A \u2014 Invalid aria-posinset="0"')
p = doc.add_paragraph()
run = p.add_run('aria-posinset must be >= 1. Found on 4 pages:')
run.font.size = Pt(9)

pages_affected = [
    'Packages: <tr role="row" aria-posinset="0" aria-setsize="20" aria-level="1">',
    'Transfers: <tr role="row" aria-posinset="0" aria-setsize="1" aria-level="1">',
    'Financials: <tr role="row" aria-posinset="0" aria-setsize="0" aria-level="1">',
    'Admin: <tr role="row" aria-posinset="0" aria-setsize="20" aria-level="1">',
]
for pa in pages_affected:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(pa)
    run.font.size = Pt(8.5)
    run.font.name = 'Consolas'

add_section_label(doc, 'Problem B \u2014 Invalid aria-activedescendant Reference')
p = doc.add_paragraph()
run = p.add_run('Packages tabstrip references aria-activedescendant="packages_tabstrip_ts_active" but this ID doesn\'t exist in the DOM.')
run.font.size = Pt(9)

add_section_label(doc, 'Problem C \u2014 Treegrid Attributes on Non-Treegrid')
p = doc.add_paragraph()
run = p.add_run('Financials grid uses aria-posinset, aria-setsize, and aria-level on <tr> within a role="grid" (these are only valid for role="treegrid").')
run.font.size = Pt(9)

add_section_label(doc, 'Fix')
add_code_block(doc, '// Remove invalid aria-posinset="0" \u2014 set to 1-based index or remove\n$(\'tr[aria-posinset="0"]\').each(function(i) {\n  $(this).attr(\'aria-posinset\', i + 1);\n});\n\n// Remove treegrid-only attributes from non-treegrid rows\n$(\'[role="grid"] tr[aria-level]\').removeAttr(\'aria-level\');\n\n// Fix or remove invalid aria-activedescendant\n$(\'#packages_tabstrip\').removeAttr(\'aria-activedescendant\');')

add_section_label(doc, 'Verification')
add_checklist(doc, [
    'axe scan shows 0 aria-valid-attr-value errors',
    'axe scan shows 0 aria-conditional-attr errors',
    'Grid keyboard navigation still functional (arrow keys, Enter to select)',
    'Screen reader correctly announces row position in grid',
    'VPAT 4.1.2 can be updated to "Supports"',
])

# ============================================================
# A11Y-005
# ============================================================
doc.add_page_break()

add_ticket_header(doc,
    'A11Y-005', 'Fix Color Contrast on Finished Grid Rows',
    'P1 - SERIOUS', '12+', '1.4.3 Contrast Minimum (Level AA)',
    '.grid-finished-row in Packages grid', '30 minutes'
)

add_section_label(doc, 'VPAT Finding')
p = doc.add_paragraph()
run = p.add_run('"Pages with grids containing icons indicating their \'type\' (Trade Sample icon, Product Package icon, Lab Sample Package, etc.) the icons are not meeting the color contrast ratio threshold"')
run.font.size = Pt(9)
run.italic = True

add_section_label(doc, 'Problem A \u2014 Text Contrast in Finished Rows')
p = doc.add_paragraph()
run = p.add_run('Text in .grid-finished-row cells has insufficient contrast:')
run.font.size = Pt(9)

contrast_details = [
    ['Foreground', '#999999 (gray)'],
    ['Background', '#f9f8f3 (off-white)'],
    ['Measured Ratio', '2.67:1'],
    ['Required', '4.5:1 for normal text (12px)'],
]
add_styled_table(doc, ['Property', 'Value'], contrast_details)

p = doc.add_paragraph()
run = p.add_run('Affected columns: Label (tag ID), Source Harvest Names, Source Package Labels, Location Name, Item Name, Product Category, Strain, Quantity, Source Production Batch, Lab Testing State, Packaged Date.')
run.font.size = Pt(9)

add_section_label(doc, 'Problem B \u2014 Icon Contrast in Grid Type Indicators')
p = doc.add_paragraph()
run = p.add_run('Grid icons indicating package type (Trade Sample, Product Package, Lab Sample Package, etc.) don\'t meet 3:1 contrast ratio for non-text elements per WCAG 1.4.11.')
run.font.size = Pt(9)

add_section_label(doc, 'Fix A \u2014 Text Contrast')
add_code_block(doc, '/* Before */\n.grid-finished-row { color: #999999; }  /* 2.67:1 \u2014 FAIL */\n\n/* After \u2014 Option 1: darker gray */\n.grid-finished-row { color: #595959; }  /* 7.01:1 \u2014 PASS */\n\n/* After \u2014 Option 2: medium gray (still visually distinct) */\n.grid-finished-row { color: #757575; }  /* 4.68:1 \u2014 PASS */')

add_section_label(doc, 'Fix B \u2014 Icon Contrast')
p = doc.add_paragraph()
run = p.add_run('Identify all grid type icons (.icon-th-large, trade sample icon, lab sample icon, etc.) and ensure icon fill/stroke colors meet 3:1 contrast ratio against #f9f8f3 background. Alternatively, add visible text labels alongside icons for redundancy.')
run.font.size = Pt(9)

add_section_label(doc, 'Verification')
add_checklist(doc, [
    'axe scan shows 0 color-contrast errors on Packages page',
    'Manual check: Finished rows visually distinct from active rows',
    'Manual check: All grid type icons meet 3:1 contrast ratio',
    'Design review/approval on new color values',
    'Check other pages with grids for similar finished-row patterns',
    'VPAT 1.4.3 can be updated to "Supports"',
])

# ============================================================
# A11Y-006
# ============================================================
add_ticket_header(doc,
    'A11Y-006', 'Fix Pointer Cancellation on Modal Buttons',
    'P1 - SERIOUS', 'TBD', '2.5.2 Pointer Cancellation (Level A)',
    'Modal-triggering buttons (application-wide)', '1 hour'
)

p = doc.add_paragraph()
run = p.add_run('\u26a0\ufe0f  Note: This issue was identified in VPAT but NOT flagged by axe scan. Requires manual investigation.')
run.font.size = Pt(9)
run.bold = True
run.font.color.rgb = SERIOUS_ORANGE

add_section_label(doc, 'VPAT Finding')
p = doc.add_paragraph()
run = p.add_run('"If a user down-clicks on a button and navigates away to up-click, the action does not perform to open the modal."')
run.font.size = Pt(9)
run.italic = True

add_section_label(doc, 'Problem')
p = doc.add_paragraph()
run = p.add_run('When a user mouse-downs on a button, then moves the cursor away before mouse-up, the action should NOT execute. The VPAT indicates this is failing for modal-opening buttons \u2014 meaning mousedown/pointerdown events are triggering actions instead of click events.')
run.font.size = Pt(9)

add_section_label(doc, 'Investigation Steps')
steps = [
    'Identify all buttons that open modals (look for mousedown or pointerdown event handlers)',
    'Test each: mouse-down on button, drag cursor off button, release \u2014 action should NOT fire',
    'Check custom button/modal components for event binding patterns',
    'Review any third-party widget event handlers (Kendo UI, Bootstrap modals)',
]
for s in steps:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(s)
    run.font.size = Pt(9)

add_section_label(doc, 'Fix')
p = doc.add_paragraph()
run = p.add_run('Ensure all interactive elements use click events (which require down+up on the same element):')
run.font.size = Pt(9)

add_code_block(doc, '// BAD \u2014 fires on mouse-down, no cancellation possible\nbutton.addEventListener(\'mousedown\', openModal);\n\n// GOOD \u2014 fires on click (down+up on same element)\nbutton.addEventListener(\'click\', openModal);')

add_code_block(doc, '// For components requiring pointerdown, implement abort-on-leave:\nbutton.addEventListener(\'pointerdown\', () => { armed = true; });\nbutton.addEventListener(\'pointerleave\', () => { armed = false; });\nbutton.addEventListener(\'pointerup\', () => { if (armed) openModal(); });')

add_section_label(doc, 'Verification')
add_checklist(doc, [
    'Manual test: Mouse-down on button, drag away, mouse-up \u2014 modal does NOT open',
    'Code audit: No mousedown/pointerdown handlers that trigger navigation or modal actions',
    'Test on touch devices: touchstart + drag away should not trigger action',
    'VPAT 2.5.2 can be updated to "Supports"',
])

# ============================================================
# A11Y-007
# ============================================================
add_ticket_header(doc,
    'A11Y-007', 'External Training Content Accessibility (Out of Scope)',
    'P2 - LOW', 'N/A', '2.4.5 Multiple Ways (Level AA)',
    'External training videos/content', 'N/A (content/UX issue)'
)

add_section_label(doc, 'VPAT Finding')
p = doc.add_paragraph()
run = p.add_run('"Lumen Expert Training/Training Videos are provided on web pages outside of the application"')
run.font.size = Pt(9)
run.italic = True

add_section_label(doc, 'Description')
p = doc.add_paragraph()
run = p.add_run('This is a content delivery issue \u2014 the application itself doesn\'t provide multiple ways to find training content. Not a same-day code fix; requires product/content strategy decisions.')
run.font.size = Pt(9)

add_section_label(doc, 'Recommendations')
recs = [
    'Add in-app search functionality that includes training content results',
    'Provide a sitemap or help index page that links to training resources',
    'Consider embedding training videos within the application with proper captions',
]
for r in recs:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(r)
    run.font.size = Pt(9)

# ============================================================
# DEPLOYMENT CHECKLIST
# ============================================================
doc.add_page_break()

add_heading_styled(doc, 'Deployment Checklist', level=1)

add_section_label(doc, 'Pre-Deploy')
add_checklist(doc, [
    'A11Y-001: role="none" on all nav <li> elements + Data Import aria-label fix',
    'A11Y-002: Move role="tablist" to <ul.k-tabstrip-items>',
    'A11Y-003: Verified resolved by A11Y-001 and A11Y-002',
    'A11Y-004: Fix aria-posinset, remove invalid aria-activedescendant, remove treegrid attrs',
    'A11Y-005: Update .grid-finished-row color + icon contrast',
    'A11Y-006: Audit and fix pointer cancellation on modal buttons',
    'Code reviewed',
])

add_section_label(doc, 'Post-Deploy Verification')
add_checklist(doc, [
    'Run axe scan on all 7 pages: Plants, Packages, Transfers, Financials, Reports, Admin, User Profile',
    'Confirm 0 critical issues, 0 serious issues',
    'Manual keyboard navigation test on main navigation',
    'Manual pointer cancellation test on modal buttons',
    'Screen reader spot check (NVDA/JAWS on Windows, VoiceOver on macOS)',
])

add_section_label(doc, 'VPAT Update')
p = doc.add_paragraph()
run = p.add_run('After all fixes verified, update conformance report:')
run.font.size = Pt(9)

add_checklist(doc, [
    '1.3.1 Info and Relationships \u2192 "Supports"',
    '4.1.2 Name, Role, Value \u2192 "Supports"',
    '1.4.3 Contrast (Minimum) \u2192 "Supports"',
    '2.5.2 Pointer Cancellation \u2192 "Supports"',
    '2.4.5 Multiple Ways \u2192 remains "Partially Supports" (external content)',
])

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/lanaholston/Desktop/Code/VPAT-Remediation-Report-Michigan-2026.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.1f} KB')
